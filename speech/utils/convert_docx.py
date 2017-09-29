from docx import Document


def convert_to_docx(title, date, content):
    document = Document()
    document.add_heading(title, 0)

    content = document.add_paragraph(content)
    return document
